# insert_calendar_docx.py
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import calendar
import sys
import os
from datetime import datetime

def make_calendar_table(doc, year, month, title=None):
    cal = calendar.Calendar(firstweekday=6)  # شنبه به عنوان ابتدای هفته (مانند tim.ir معمولاً شنبه)
    month_days = cal.monthdayscalendar(year, month)  # لیست هفته‌ها

    if title is None:
        title = f"{calendar.month_name[month]} {year}"

    # صفحه جدید
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(16)
    run.bold = True
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # ایجاد جدول: 7 ستون (ش-ی-د-س-چ-پ-ج) و به اندازه ردیف‌های لازم + 1 برای header
    rows = len(month_days) + 1
    table = doc.add_table(rows=rows, cols=7)
    table.style = 'Table Grid'

    # هدر روزهای هفته (اینجا به فارسی)
    weekdays = ['ش', 'ی', 'د', 'س', 'چ', 'پ', 'ج']  # شنبه..جمعه
    hdr_cells = table.rows[0].cells
    for i, wd in enumerate(weekdays):
        cell = hdr_cells[i]
        cell.text = wd
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.bold = True

    # پر کردن روزها
    for r, week in enumerate(month_days, start=1):
        row_cells = table.rows[r].cells
        for c, day in enumerate(week):
            text = '' if day == 0 else str(day)
            cell = row_cells[c]
            cell.text = text
            # تنظیم فونت و تراز
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    # یک پاراگراف توضیحی پایین جدول تا کاربر بتواند توضیحات اضافه کند
    doc.add_paragraph("\nتوضیحات: ")

def insert_calendar_into_docx(path_in, path_out, year, month, title=None):
    if not os.path.exists(path_in):
        # اگر فایل وجود ندارد، یک فایل جدید بساز
        doc = Document()
        doc.add_paragraph("تقویم اضافه شده: ")
    else:
        doc = Document(path_in)

    make_calendar_table(doc, year, month, title)
    doc.save(path_out)
    print(f"Saved: {path_out}")

if __name__ == "__main__":
    # مثال: python insert_calendar_docx.py input.docx output.docx 2025 11
    if len(sys.argv) < 5:
        print("Usage: python insert_calendar_docx.py <input.docx or new> <output.docx> <year> <month>")
        sys.exit(1)
    in_path = sys.argv[1]
    out_path = sys.argv[2]
    year = int(sys.argv[3])
    month = int(sys.argv[4])
    insert_calendar_into_docx(in_path if in_path.lower() != "new" else None, out_path, year, month)
